import hashlib
import io
import os
from typing import Dict, Any, Tuple
from fastapi import HTTPException, status
import fitz  # PyMuPDF
import docx

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

def compute_sha256(file_bytes: bytes) -> str:
    """Compute SHA-256 hash of file content for deduplication and caching."""
    return hashlib.sha256(file_bytes).hexdigest()

def validate_file(filename: str, file_bytes: bytes) -> str:
    """Validate file extension and size. Returns file extension (lowercase)."""
    _, ext = os.path.splitext(filename.lower())
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file format '{ext}'. Allowed formats: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds maximum allowed limit of {MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB"
        )
    if len(file_bytes) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file is empty (0 bytes)."
        )
    return ext.lstrip(".")

def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extract text and page count from a PDF file in-memory using PyMuPDF (fitz).
    Handles text encoding, multi-page flows, and layout blocks.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = len(doc)
        text_parts = []
        
        for page_idx in range(page_count):
            page = doc[page_idx]
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text.strip())
        
        doc.close()
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="PDF appears to be scanned or contains only images with no extractable text. Please upload a searchable text PDF."
            )
            
        return full_text, page_count
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to extract text from PDF: {str(e)}"
        )

def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extract text from a Word document (.docx) in-memory using python-docx.
    Extracts paragraphs, bullet items, and table cells.
    """
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        text_parts = []

        # Extract paragraphs
        for p in doc.paragraphs:
            clean_p = p.text.strip()
            if clean_p:
                text_parts.append(clean_p)

        # Extract table content (e.g. skills or experience matrices)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        full_text = "\n\n".join(text_parts)
        if not full_text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="DOCX file contains no extractable text content."
            )

        # Estimate page count (standard assumption ~450 words per page)
        words = len(full_text.split())
        estimated_pages = max(1, round(words / 450))
        return full_text, estimated_pages
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to extract text from DOCX: {str(e)}"
        )

def parse_resume_document(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Primary ingestion pipeline:
    1. Validates format and size.
    2. Calculates SHA-256 hash.
    3. Extracts text and page metadata.
    """
    file_type = validate_file(filename, file_bytes)
    file_hash = compute_sha256(file_bytes)

    if file_type == "pdf":
        raw_text, page_count = extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        raw_text, page_count = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported format: {file_type}"
        )

    words = raw_text.split()
    word_count = len(words)
    char_count = len(raw_text)

    return {
        "filename": filename,
        "file_type": file_type,
        "file_hash": file_hash,
        "raw_text": raw_text,
        "page_count": page_count,
        "word_count": word_count,
        "char_count": char_count,
        "file_size_bytes": len(file_bytes)
    }
