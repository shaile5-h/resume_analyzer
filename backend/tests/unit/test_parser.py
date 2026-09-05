import os
import pytest
import fitz
import docx
from fastapi import HTTPException
from app.services.parser import (
    validate_file,
    compute_sha256,
    extract_text_from_pdf,
    extract_text_from_docx,
    parse_resume_document
)

def test_compute_sha256():
    data = b"Resume text content"
    hash1 = compute_sha256(data)
    hash2 = compute_sha256(data)
    assert hash1 == hash2
    assert len(hash1) == 64

def test_validate_file_valid_extensions():
    pdf_ext = validate_file("resume.PDF", b"some bytes")
    assert pdf_ext == "pdf"
    
    docx_ext = validate_file("my_cv.docx", b"some bytes")
    assert docx_ext == "docx"

def test_validate_file_unsupported_extension():
    with pytest.raises(HTTPException) as exc:
        validate_file("resume.txt", b"some bytes")
    assert exc.value.status_code == 400
    assert "Unsupported file format" in exc.value.detail

def test_validate_file_empty_content():
    with pytest.raises(HTTPException) as exc:
        validate_file("resume.pdf", b"")
    assert exc.value.status_code == 400
    assert "empty" in exc.value.detail

def test_validate_file_oversized():
    oversized_bytes = b"x" * (11 * 1024 * 1024) # 11 MB
    with pytest.raises(HTTPException) as exc:
        validate_file("resume.pdf", oversized_bytes)
    assert exc.value.status_code == 413
    assert "maximum allowed limit" in exc.value.detail

def test_extract_text_from_pdf():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Senior Python Engineer\nExperience in FastAPI and Docker.")
    pdf_bytes = doc.tobytes()
    doc.close()

    text, pages = extract_text_from_pdf(pdf_bytes)
    assert pages == 1
    assert "Senior Python Engineer" in text
    assert "FastAPI" in text

def test_extract_text_from_docx(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("DevOps Engineer Resume")
    doc.add_paragraph("AWS, Kubernetes, Terraform")
    
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skill"
    table.rows[0].cells[1].text = "Proficiency"

    temp_file = tmp_path / "test.docx"
    doc.save(str(temp_file))
    
    with open(temp_file, "rb") as f:
        docx_bytes = f.read()

    text, pages = extract_text_from_docx(docx_bytes)
    assert pages >= 1
    assert "DevOps Engineer Resume" in text
    assert "Kubernetes" in text
    assert "Skill | Proficiency" in text

def test_parse_resume_document():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "John Doe\nSoftware Engineer")
    pdf_bytes = doc.tobytes()
    doc.close()

    result = parse_resume_document(pdf_bytes, "test.pdf")
    assert result["filename"] == "test.pdf"
    assert result["file_type"] == "pdf"
    assert result["page_count"] == 1
    assert result["word_count"] >= 3
    assert len(result["file_hash"]) == 64
    assert "John Doe" in result["raw_text"]
