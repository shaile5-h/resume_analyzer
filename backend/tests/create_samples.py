"""
Utility to generate 6 realistic test resumes in PDF and DOCX formats for integration testing.
"""

import os
import fitz  # PyMuPDF
import docx
from docx.shared import Pt, Inches, RGBColor

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEST_RESUMES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_resumes")
ROOT_RESUMES_DIR = os.path.join(os.path.dirname(BASE_DIR), "sample_resumes")

for d in [TEST_RESUMES_DIR, ROOT_RESUMES_DIR]:
    os.makedirs(d, exist_ok=True)

def build_pdf(filename: str, name: str, contact: str, sections: list):
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 50
    margin_x = 50
    usable_width = 495

    page.insert_text((margin_x, y), name, fontsize=20, fontname="helv", color=(0.1, 0.15, 0.3))
    y += 24
    page.insert_text((margin_x, y), contact, fontsize=9.5, fontname="helv", color=(0.3, 0.35, 0.4))
    y += 18
    page.draw_line((margin_x, y), (margin_x + usable_width, y), color=(0.8, 0.85, 0.9), width=1)
    y += 16

    for heading, items in sections:
        y += 8
        page.insert_text((margin_x, y), heading.upper(), fontsize=11, fontname="helv", color=(0.2, 0.25, 0.5))
        y += 4
        page.draw_line((margin_x, y), (margin_x + usable_width, y), color=(0.7, 0.75, 0.85), width=0.8)
        y += 14

        for item in items:
            if item.startswith("SUBTITLE:"):
                subtitle = item.replace("SUBTITLE:", "").strip()
                page.insert_text((margin_x, y), subtitle, fontsize=10.5, fontname="helv", color=(0.15, 0.15, 0.2))
                y += 14
            elif item.startswith("BULLET:"):
                bullet = f"•   {item.replace('BULLET:', '').strip()}"
                words = bullet.split()
                line = ""
                for w in words:
                    test_l = f"{line} {w}".strip()
                    if len(test_l) * 5.2 > usable_width:
                        page.insert_text((margin_x + 10, y), line, fontsize=9, fontname="helv", color=(0.25, 0.25, 0.25))
                        y += 13
                        line = f"    {w}"
                    else:
                        line = test_l
                if line:
                    page.insert_text((margin_x + 10, y), line, fontsize=9, fontname="helv", color=(0.25, 0.25, 0.25))
                    y += 13
            else:
                words = item.split()
                line = ""
                for w in words:
                    test_l = f"{line} {w}".strip()
                    if len(test_l) * 5.2 > usable_width:
                        page.insert_text((margin_x, y), line, fontsize=9, fontname="helv", color=(0.25, 0.25, 0.25))
                        y += 13
                        line = w
                    else:
                        line = test_l
                if line:
                    page.insert_text((margin_x, y), line, fontsize=9, fontname="helv", color=(0.25, 0.25, 0.25))
                    y += 13
                y += 4

    # Save to both sample_resumes/ and backend/tests/sample_resumes/
    for folder in [TEST_RESUMES_DIR, ROOT_RESUMES_DIR]:
        path = os.path.join(folder, filename)
        doc.save(path)
    doc.close()
    print(f"Generated PDF: {filename}")

def build_docx(filename: str, name: str, contact: str, sections: list):
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    p_name = doc.add_paragraph()
    r_name = p_name.add_run(name)
    r_name.font.size = Pt(20)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(30, 41, 59)
    p_name.paragraph_format.space_after = Pt(2)

    p_contact = doc.add_paragraph()
    r_contact = p_contact.add_run(contact)
    r_contact.font.size = Pt(9.5)
    r_contact.font.color.rgb = RGBColor(100, 116, 139)
    p_contact.paragraph_format.space_after = Pt(12)

    for heading, items in sections:
        p_h = doc.add_paragraph()
        r_h = p_h.add_run(heading.upper())
        r_h.font.size = Pt(11)
        r_h.font.bold = True
        r_h.font.color.rgb = RGBColor(67, 56, 202)
        p_h.paragraph_format.space_before = Pt(8)
        p_h.paragraph_format.space_after = Pt(4)

        for item in items:
            if item.startswith("SUBTITLE:"):
                p_sub = doc.add_paragraph()
                r_sub = p_sub.add_run(item.replace("SUBTITLE:", "").strip())
                r_sub.font.size = Pt(10)
                r_sub.font.bold = True
                r_sub.font.color.rgb = RGBColor(51, 65, 85)
                p_sub.paragraph_format.space_after = Pt(2)
            elif item.startswith("BULLET:"):
                p_b = doc.add_paragraph(style="List Bullet")
                r_b = p_b.add_run(item.replace("BULLET:", "").strip())
                r_b.font.size = Pt(9.5)
                r_b.font.color.rgb = RGBColor(71, 85, 105)
                p_b.paragraph_format.space_after = Pt(2)
            else:
                p_p = doc.add_paragraph()
                r_p = p_p.add_run(item)
                r_p.font.size = Pt(9.5)
                r_p.font.color.rgb = RGBColor(71, 85, 105)
                p_p.paragraph_format.space_after = Pt(4)

    for folder in [TEST_RESUMES_DIR, ROOT_RESUMES_DIR]:
        path = os.path.join(folder, filename)
        doc.save(path)
    print(f"Generated DOCX: {filename}")

def generate_samples():
    # 1. Senior Full Stack Engineer (PDF)
    build_pdf(
        "1_Senior_FullStack_Engineer.pdf",
        "Alexander Wright",
        "alex.wright@techmail.com | (555) 234-5678 | San Francisco, CA | linkedin.com/in/alexanderwright | github.com/awright-dev",
        [
            ("Professional Summary", ["Results-driven Senior Full-Stack Engineer with 7+ years of experience architecting resilient microservices and responsive web applications. Proven track record of spearheading cloud migrations that scaled traffic by 300% while cutting infrastructure costs by $140k annually."]),
            ("Technical Skills", ["Languages: Python, TypeScript, JavaScript, Go, SQL, HTML5, CSS3\nFrameworks: FastAPI, Django, React.js, Next.js, Node.js, Tailwind CSS\nCloud & DevOps: AWS (ECS, S3, RDS, Lambda), Docker, Kubernetes, Terraform, CI/CD\nDatabases: PostgreSQL, Redis, SQLite, MongoDB"]),
            ("Work Experience", [
                "SUBTITLE: Senior Full-Stack Engineer | CloudScale Technologies (2021 – Present)",
                "BULLET: Architected and engineered distributed microservices handling 2.5M daily active requests using FastAPI and PostgreSQL with 99.99% uptime.",
                "BULLET: Optimized database query performance and introduced multi-tier Redis caching, cutting average API response latency by 45%.",
                "BULLET: Spearheaded frontend migration to React and TypeScript, boosting core web vitals and reducing page load duration by 35%.",
                "BULLET: Mentored a team of 6 engineers, establishing automated CI/CD pipelines and elevating unit test coverage to 92%.",
                "SUBTITLE: Software Engineer | Apex Solutions Inc. (2018 – 2021)",
                "BULLET: Developed 15+ RESTful endpoints in Python and React supporting a high-volume B2B e-commerce platform generating $18M in ARR.",
                "BULLET: Automated deployment processes with Docker and GitHub Actions, reducing release cycles from 4 hours to 15 minutes."
            ]),
            ("Education", [
                "SUBTITLE: Bachelor of Science in Computer Science | University of California, Berkeley (2014 – 2018)",
                "BULLET: Graduated Magna Cum Laude, GPA: 3.85 / 4.0"
            ])
        ]
    )

    # 2. Data Scientist & ML Engineer (PDF)
    build_pdf(
        "2_Data_Scientist_ML_Engineer.pdf",
        "Dr. Sophia Chen",
        "sophia.chen@mlresearch.io | (555) 789-0123 | New York, NY | linkedin.com/in/sophiachen-phd | github.com/sophiachen-ml",
        [
            ("Professional Summary", ["Lead Machine Learning Engineer with 6+ years of expertise delivering production deep learning and predictive models. Specialized in NLP, LLM fine-tuning, RAG pipelines, and scalable feature engineering on distributed GPU clusters."]),
            ("Technical Skills", ["Core: Python, R, C++, PyTorch, TensorFlow, Scikit-learn, HuggingFace, LangChain\nData & Cloud: Spark, Kafka, Snowflake, GCP Vertex AI, AWS SageMaker, Docker, MLflow\nTechniques: Large Language Models, Transformer Architectures, Semantic Search, Time-Series Forecasting"]),
            ("Work Experience", [
                "SUBTITLE: Lead Machine Learning Scientist | Omniscient AI (2022 – Present)",
                "BULLET: Engineered enterprise RAG pipeline utilizing embedding models and vector databases, boosting document retrieval precision by 42% for 85k enterprise users.",
                "BULLET: Fine-tuned open-source LLMs on proprietary financial corpora, reducing GPU inference overhead by 28% while improving accuracy by 15%.",
                "BULLET: Deployed real-time fraud detection pipeline processing $50M in daily transactions with under 40ms inference latency.",
                "SUBTITLE: Data Scientist | Horizon Analytics (2019 – 2022)",
                "BULLET: Built predictive customer churn model that retained $2.4M in annual subscription revenue across 120 enterprise accounts.",
                "BULLET: Spearheaded A/B testing infrastructure evaluating recommendation algorithms, yielding an 18% lift in user engagement."
            ]),
            ("Education", [
                "SUBTITLE: Ph.D. in Computer Science (Machine Learning Focus) | Columbia University (2015 – 2019)",
                "SUBTITLE: B.S. in Applied Mathematics | MIT (2011 – 2015)"
            ])
        ]
    )

    # 3. DevOps & Cloud Infrastructure Architect (DOCX)
    build_docx(
        "3_DevOps_Cloud_Architect.docx",
        "Marcus Sterling",
        "marcus.sterling@devopscloud.net | (555) 456-7890 | Austin, TX | linkedin.com/in/marcussterling | github.com/msterling-ops",
        [
            ("Professional Summary", ["Staff Cloud & DevOps Architect with 8+ years of experience leading multi-cloud infrastructure transformations, zero-trust security postures, and Kubernetes cluster orchestration. Proven record of achieving 99.999% system availability and saving $320k in annual compute expenses."]),
            ("Technical Skills", ["Cloud Platforms: AWS (EKS, VPC, CloudFront, IAM), Google Cloud Platform (GKE), Azure\nInfrastructure as Code: Terraform, Terragrunt, Ansible, CloudFormation\nCI/CD & Containers: Kubernetes, Docker, Helm, ArgoCD, GitLab CI, GitHub Actions, Jenkins\nObservability: Prometheus, Grafana, Datadog, OpenTelemetry, ELK Stack"]),
            ("Work Experience", [
                "SUBTITLE: Principal DevOps Architect | FinTech Velocity (2021 – Present)",
                "BULLET: Architected multi-region Kubernetes cluster deployment on AWS handling 15,000 requests per second with automated failover and zero data loss.",
                "BULLET: Consolidated 40+ legacy services into automated Terraform modules, reducing new environment provisioning time from 5 days to 20 minutes.",
                "BULLET: Introduced GitOps workflow using ArgoCD and Helm, decreasing deployment failure rates by 65% across 24 engineering teams.",
                "BULLET: Overhauled AWS cost governance and rightsized Spot instance pools, reducing annual cloud compute bill by $320,000.",
                "SUBTITLE: Cloud Infrastructure Engineer | ScaleCloud Partners (2017 – 2021)",
                "BULLET: Spearheaded migration of on-premise datacenter workloads to AWS VPC, finishing 2 months ahead of schedule.",
                "BULLET: Configured Prometheus and Grafana alerts across 200+ nodes, cutting Mean Time to Detect (MTTD) incidents by 50%."
            ]),
            ("Education", [
                "SUBTITLE: Bachelor of Science in Information Systems | University of Texas at Austin (2013 – 2017)"
            ]),
            ("Certifications", [
                "AWS Certified Solutions Architect – Professional\nCertified Kubernetes Administrator (CKA)\nHashiCorp Certified: Terraform Associate"
            ])
        ]
    )

    # 4. Junior Frontend Web Developer (DOCX)
    build_docx(
        "4_Junior_Frontend_Developer.docx",
        "Emily Johnson",
        "emily.johnson.dev@gmail.com | (555) 345-6789 | Seattle, WA | linkedin.com/in/emilyj-dev | github.com/emilyj-frontend",
        [
            ("Professional Summary", ["Enthusiastic Junior Frontend Developer with 1.5 years of experience building modern, accessible, and responsive user interfaces using React, TypeScript, and Tailwind CSS. Passionate about clean UI design and optimizing user experience."]),
            ("Technical Skills", ["Languages: JavaScript (ES6+), TypeScript, HTML5, CSS3, Python\nFrameworks: React.js, Vite, Tailwind CSS, Bootstrap\nTools: Git, GitHub, VS Code, Figma, REST APIs, Jest"]),
            ("Work Experience", [
                "SUBTITLE: Associate Frontend Developer | PixelCraft Studio (2023 – Present)",
                "BULLET: Developed responsive components using React.js and Tailwind CSS for 4 client web applications.",
                "BULLET: Assisted in integrating RESTful APIs to display dynamic product catalogues and user profile details.",
                "BULLET: Improved mobile accessibility compliance from 78% to 95% on Lighthouse audits."
            ]),
            ("Projects", [
                "SUBTITLE: AI Recipe Finder Web App (Personal Project | 2023)",
                "BULLET: Built a single-page React app connecting to open recipes API, attracting 1,200 active monthly users.",
                "BULLET: Implemented local storage caching to maintain favorite bookmarks offline."
            ]),
            ("Education", [
                "SUBTITLE: Bachelor of Science in Web Design & Development | University of Washington (2019 – 2023)"
            ])
        ]
    )

    # 5. Technical Product Manager (PDF)
    build_pdf(
        "5_Product_Manager_Tech.pdf",
        "David K. Martinez",
        "david.martinez@pmleadership.com | (555) 901-2345 | Chicago, IL | linkedin.com/in/davidkmartinez",
        [
            ("Professional Summary", ["Customer-centric Technical Product Manager with 5+ years of experience leading cross-functional engineering and design teams. Proven expertise in taking B2B SaaS platforms from concept to product-market fit, generating over $8.5M in incremental annual revenue."]),
            ("Core Competencies", ["Product Strategy, Roadmap Prioritization, Agile & Scrum, User Research, Wireframing, SQL, Data Analytics, Mixpanel, Jira, Go-to-Market Strategy, Stakeholder Management"]),
            ("Work Experience", [
                "SUBTITLE: Senior Product Manager | Elevate Software (2021 – Present)",
                "BULLET: Spearheaded end-to-end product development for enterprise analytics portal, driving $4.2M in net-new ARR within 12 months of launch.",
                "BULLET: Conducted 60+ user discovery interviews and leveraged Mixpanel behavioral metrics to optimize onboarding, improving user activation by 32%.",
                "BULLET: Managed backlog and sprint ceremonies for two 8-person engineering squads, delivering 95% of committed roadmap features on schedule.",
                "SUBTITLE: Product Manager | MetricFlow Systems (2018 – 2021)",
                "BULLET: Launched self-serve payment integration feature that lowered churn by 18% and accelerated customer time-to-value by 40%.",
                "BULLET: Partnered with sales and marketing teams to establish customer feedback loops, boosting NPS score from 42 to 68."
            ]),
            ("Education", [
                "SUBTITLE: Master of Business Administration (MBA) | Northwestern University, Kellogg (2016 – 2018)",
                "SUBTITLE: B.S. in Industrial Engineering | University of Illinois Urbana-Champaign (2012 – 2016)"
            ])
        ]
    )

    # 6. Low-score edge case resume (PDF)
    build_pdf(
        "6_Needs_Improvement_EdgeCase.pdf",
        "John Doe",
        "Location: Dallas, Texas (No email or phone listed)",
        [
            ("Things I Have Worked On", ["I have been doing coding for many years in different companies and projects. I like solving computer problems and writing code for websites and systems."]),
            ("My Past Roles", [
                "SUBTITLE: Developer at Small Tech Firm (Few years ago)",
                "BULLET: Worked on computer bugs and assisted senior teammates with code.",
                "BULLET: Did some website modifications and made pages look better.",
                "BULLET: Attended company meetings and discussed project goals."
            ]),
            ("What I Know", ["Computers, internet, programming languages, databases, problem solving."]),
            ("Schooling", ["College degree in general technology studies."])
        ]
    )

if __name__ == "__main__":
    generate_samples()
